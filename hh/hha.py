import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile
import webbrowser
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime

APP_TITLE = "برنامج مذكرات أساتذة المتوسط"
DATA_FILE = "teacher_notes_data.json"
SCHOOL_NAME = "متوسطة النجاح"
SCHOOL_CITY = "الجزائر"

BG = "#0f172a"
PANEL = "#111827"
CARD = "#1f2937"
GOLD = "#d4af37"
TEXT = "#f9fafb"
MUTED = "#cbd5e1"
ACCENT = "#2563eb"
SUCCESS = "#16a34a"
WARN = "#f59e0b"
DANGER = "#dc2626"
INPUT_BG = "#f8fafc"
INPUT_FG = "black"

TEMPLATES = {
    "اللغة العربية": {
        "الكفاءة الختامية": "يقرأ المتعلم النص ويفهم مضمونه ويستخرج أفكاره ويوظف مكتسباته اللغوية في وضعيات مختلفة.",
        "مؤشر الكفاءة": "أن يحدد الفكرة العامة والأفكار الأساسية، وأن يجيب عن أسئلة الفهم، وأن يوظف الرصيد اللغوي توظيفا سليما.",
        "الوسائل": "الكتاب المدرسي، السبورة، كراس النشاطات، بطاقات.",
        "وضعية الانطلاق": "مراجعة سريعة لمكتسبات الحصة السابقة وطرح أسئلة تمهيدية مرتبطة بعنوان الدرس.",
        "سير الحصة": "تمهيد - قراءة صامتة أو جهرية - شرح المفردات - مناقشة الفهم - استخراج الأفكار - استثمار لغوي - تطبيق.",
        "التقويم": "تقويم شفهي وكتابي من خلال أسئلة مباشرة وتمارين تطبيقية قصيرة.",
        "الوضعية الإدماجية": "يوظف المتعلم التعلمات الجديدة في فقرة أو نشاط إدماجي مناسب.",
        "الملاحظات": "مراعاة الفروق الفردية وتشجيع المتعلمين على المشاركة."
    },
    "الرياضيات": {
        "الكفاءة الختامية": "يحل المتعلم وضعيات مشكلة بتوظيف المفاهيم والقواعد الرياضية المناسبة.",
        "مؤشر الكفاءة": "أن ينجز العمليات بدقة، ويستعمل القاعدة المناسبة، ويبرر خطوات الحل.",
        "الوسائل": "الكتاب المدرسي، السبورة، أدوات هندسية، أوراق عمل.",
        "وضعية الانطلاق": "طرح وضعية إشكالية قصيرة مرتبطة بحياة المتعلم تقوده لاكتشاف المفهوم الجديد.",
        "سير الحصة": "عرض الوضعية - البحث الفردي أو الجماعي - المناقشة - استخلاص القاعدة - تطبيقات تدريجية.",
        "التقويم": "تقويم مرحلي أثناء الإنجاز ثم تقويم نهائي بتمارين تطبيقية.",
        "الوضعية الإدماجية": "حل وضعية مركبة تستدعي توظيف التعلمات المكتسبة في الحصة.",
        "الملاحظات": "التركيز على منهجية الحل وتنظيم الكتابة الرياضية."
    },
    "العلوم الطبيعية": {
        "الكفاءة الختامية": "يفسر المتعلم الظواهر الطبيعية بالاعتماد على الملاحظة والتجريب والاستنتاج العلمي.",
        "مؤشر الكفاءة": "أن يصف الظاهرة، ويحلل النتائج، ويستنتج الخلاصة العلمية الصحيحة.",
        "الوسائل": "الكتاب المدرسي، صور، مجسمات، وثائق علمية، مخبر إن توفر.",
        "وضعية الانطلاق": "عرض صورة أو وضعية علمية تثير تساؤلات التلاميذ حول الظاهرة المدروسة.",
        "سير الحصة": "ملاحظة - فرضيات - نشاط أو تجربة - تحليل - استنتاج - تثبيت.",
        "التقويم": "تقويم عبر أسئلة تحليلية وتوظيف النتائج في وضعيات بسيطة.",
        "الوضعية الإدماجية": "إنجاز نشاط يربط بين المفاهيم العلمية المكتسبة وواقع المتعلم.",
        "الملاحظات": "تحفيز المتعلمين على التعبير العلمي الدقيق."
    },
    "الفيزياء": {
        "الكفاءة الختامية": "يوظف المتعلم المفاهيم الفيزيائية في تفسير الظواهر وحل وضعيات بسيطة.",
        "مؤشر الكفاءة": "أن يلاحظ ويقيس ويستنتج العلاقة الفيزيائية الصحيحة.",
        "الوسائل": "أجهزة مخبرية، الكتاب المدرسي، السبورة، بطاقات نشاط.",
        "وضعية الانطلاق": "مناقشة ظاهرة من المحيط اليومي تقود إلى التساؤل الفيزيائي.",
        "سير الحصة": "تقديم الوضعية - فرضيات - تجربة أو نشاط - قراءة النتائج - استخلاص القانون أو القاعدة.",
        "التقويم": "أسئلة مباشرة وتمارين قصيرة على المفهوم أو القانون المدروس.",
        "الوضعية الإدماجية": "حل وضعية توظف المكتسبات الفيزيائية بطريقة عملية.",
        "الملاحظات": "الاهتمام بالدقة في القياس والرموز والوحدات."
    },
    "التاريخ والجغرافيا": {
        "الكفاءة الختامية": "يستثمر المتعلم المعارف التاريخية والجغرافية لفهم الأحداث والظواهر وتحليلها.",
        "مؤشر الكفاءة": "أن يحدد الزمان والمكان، ويستخرج المعلومات، ويحلل الوثائق والخرائط.",
        "الوسائل": "الكتاب المدرسي، خرائط، وثائق، صور، سبورة.",
        "وضعية الانطلاق": "استحضار معارف سابقة وطرح أسئلة تمهيدية حول الحدث أو الظاهرة.",
        "سير الحصة": "قراءة الوثائق - تحليل - مناقشة - تركيب - تثبيت.",
        "التقويم": "تقويم عبر أسئلة فهم وتحليل ووضعيات قصيرة.",
        "الوضعية الإدماجية": "إنجاز نشاط تركيبي أو تعليق موجز على وثيقة أو خريطة.",
        "الملاحظات": "التركيز على الربط بين السبب والنتيجة ودقة المصطلحات."
    },
    "اللغة الفرنسية": {
        "الكفاءة الختامية": "Communiquer correctement à l'oral et à l'écrit dans des situations adaptées au niveau des apprenants.",
        "مؤشر الكفاءة": "Comprendre, reformuler, produire des phrases correctes et réinvestir le lexique étudié.",
        "الوسائل": "Livre scolaire, tableau, images, fiches.",
        "وضعية الانطلاق": "Échange oral simple autour du thème de la leçon.",
        "سير الحصة": "Observation - compréhension - explication - entraînement - production.",
        "التقويم": "Questions orales et petite activité écrite.",
        "الوضعية الإدماجية": "Production courte pour réinvestir les acquis.",
        "الملاحظات": "Favoriser la participation des élèves et corriger progressivement."
    },
    "اللغة الإنجليزية": {
        "الكفاءة الختامية": "The learner communicates in simple oral and written situations using appropriate language.",
        "مؤشر الكفاءة": "Identify key vocabulary, understand instructions, and produce simple correct sentences.",
        "الوسائل": "Textbook, board, flashcards, worksheets.",
        "وضعية الانطلاق": "Warm-up questions related to the lesson topic.",
        "سير الحصة": "Presentation - practice - interaction - production - feedback.",
        "التقويم": "Short oral questions and written tasks.",
        "الوضعية الإدماجية": "Small integrated task using the target language.",
        "الملاحظات": "Encourage learners to speak and participate confidently."
    },
    "التربية الإسلامية": {
        "الكفاءة الختامية": "يفهم المتعلم القيم الإسلامية ويوظفها في سلوكه اليومي.",
        "مؤشر الكفاءة": "أن يستوعب المفهوم الشرعي أو الخلقي وأن يربطه بسلوكه العملي.",
        "الوسائل": "الكتاب المدرسي، السبورة، نصوص شرعية، بطاقات.",
        "وضعية الانطلاق": "طرح سؤال أو موقف حياتي مرتبط بموضوع الدرس.",
        "سير الحصة": "تمهيد - قراءة النص - شرح - استنباط الأحكام أو القيم - تطبيق.",
        "التقويم": "أسئلة شفهية وكتابية حول المفاهيم والقيم.",
        "الوضعية الإدماجية": "وضعية تربوية يوظف فيها المتعلم قيمة أو حكما شرعيا.",
        "الملاحظات": "التركيز على الجانب السلوكي والتربوي."
    },
    "الإعلام الآلي": {
        "الكفاءة الختامية": "يوظف المتعلم مهارات الإعلام الآلي لإنجاز مهام رقمية بسيطة ومنظمة.",
        "مؤشر الكفاءة": "أن يستعمل الأدوات الرقمية الأساسية وينفذ التعليمات بدقة.",
        "الوسائل": "حاسوب، جهاز عرض، برمجيات تعليمية، سبورة.",
        "وضعية الانطلاق": "عرض مهمة رقمية أو مشكلة بسيطة مرتبطة بالحاسوب.",
        "سير الحصة": "شرح - تطبيق موجه - إنجاز فردي - تصحيح ومناقشة.",
        "التقويم": "تقويم عملي من خلال تنفيذ مهمة رقمية قصيرة.",
        "الوضعية الإدماجية": "إنجاز نشاط تطبيقي يستثمر المهارات الرقمية المكتسبة.",
        "الملاحظات": "مرافقة المتعلمين أثناء الإنجاز ومراعاة الفروق في الإيقاع."
    }
}


class NotesApp:
    def __init__(self, root):
        self.root = root
        self.root.title(APP_TITLE)
        self.root.geometry("1380x860")
        self.root.configure(bg=BG)

        self.notes = []
        self.current_index = None
        self.entries = {}
        self.texts = {}
        self.search_var = tk.StringVar()
        self.count_var = tk.StringVar(value="0")
        self.subject_var = tk.StringVar(value="-")
        self.level_var = tk.StringVar(value="-")

        self.load_data()
        self.build_ui()
        self.refresh_table()

    def make_button(self, parent, text, command, bg, fg="white"):
        return tk.Button(
            parent,
            text=text,
            command=command,
            bg=bg,
            fg=fg,
            font=("Arial", 11, "bold"),
            padx=12,
            pady=8,
            relief="raised",
            bd=2,
            cursor="hand2",
        )

    def build_ui(self):
        header = tk.Frame(self.root, bg=BG, pady=12)
        header.pack(fill="x")
        tk.Label(header, text="📘 برنامج مذكرات أساتذة المتوسط", font=("Arial", 24, "bold"), bg=BG, fg=GOLD).pack()
        tk.Label(header, text="واجهة احترافية مع قوالب جاهزة حسب المادة وحفظ وأرشفة للمذكرات", font=("Arial", 11), bg=BG, fg=MUTED).pack(pady=4)

        stats = tk.Frame(self.root, bg=BG)
        stats.pack(fill="x", padx=10, pady=(0, 10))
        for i, (title, var) in enumerate([
            ("عدد المذكرات", self.count_var),
            ("المادة الحالية", self.subject_var),
            ("المستوى الحالي", self.level_var),
        ]):
            box = tk.Frame(stats, bg=CARD, padx=16, pady=12, highlightbackground=GOLD, highlightthickness=1)
            box.grid(row=0, column=i, padx=4, sticky="nsew")
            stats.grid_columnconfigure(i, weight=1)
            tk.Label(box, text=title, bg=CARD, fg=MUTED, font=("Arial", 11, "bold")).pack()
            tk.Label(box, textvariable=var, bg=CARD, fg=TEXT, font=("Arial", 16, "bold")).pack(pady=4)

        container = tk.Frame(self.root, bg=BG)
        container.pack(fill="both", expand=True, padx=10, pady=10)

        left = tk.LabelFrame(container, text="بيانات المذكرة", font=("Arial", 12, "bold"), bg=PANEL, fg=GOLD, padx=10, pady=10)
        left.pack(side="right", fill="y", padx=6)

        right = tk.Frame(container, bg=BG)
        right.pack(side="left", fill="both", expand=True, padx=6)

        labels = [
            "اسم الأستاذ", "المادة", "المستوى", "المقطع", "الوحدة", "عنوان الدرس",
            "الأسبوع", "الحصة", "المدة", "الكفاءة الختامية", "مؤشر الكفاءة", "الوسائل"
        ]

        for i, label in enumerate(labels):
            tk.Label(left, text=label, bg=PANEL, fg=TEXT, font=("Arial", 11, "bold")).grid(row=i, column=0, sticky="e", padx=5, pady=4)
            if label == "المستوى":
                widget = ttk.Combobox(left, values=["الأولى متوسط", "الثانية متوسط", "الثالثة متوسط", "الرابعة متوسط"], state="readonly", justify="right", width=35)
                widget.set("الأولى متوسط")
                widget.bind("<<ComboboxSelected>>", lambda e: self.update_current_labels())
            elif label == "المادة":
                widget = ttk.Combobox(left, values=list(TEMPLATES.keys()), state="readonly", justify="right", width=35)
                widget.set("اللغة العربية")
                widget.bind("<<ComboboxSelected>>", lambda e: self.apply_template())
            else:
                widget = tk.Entry(left, justify="right", width=38, font=("Arial", 11), bg=INPUT_BG, fg=INPUT_FG)
            widget.grid(row=i, column=1, padx=5, pady=4)
            self.entries[label] = widget

        multi_fields = [
            ("وضعية الانطلاق", 4),
            ("سير الحصة", 8),
            ("التقويم", 4),
            ("الوضعية الإدماجية", 4),
            ("الملاحظات", 3),
        ]

        start_row = len(labels)
        for idx, (label, height) in enumerate(multi_fields):
            tk.Label(left, text=label, bg=PANEL, fg=TEXT, font=("Arial", 11, "bold")).grid(row=start_row + idx * 2, column=0, sticky="ne", padx=5, pady=4)
            text = tk.Text(left, width=38, height=height, font=("Arial", 11), wrap="word", bg=INPUT_BG, fg=INPUT_FG)
            text.grid(row=start_row + idx * 2, column=1, padx=5, pady=4)
            self.texts[label] = text

        btn_frame = tk.Frame(left, bg=PANEL)
        btn_frame.grid(row=start_row + len(multi_fields) * 2 + 1, column=0, columnspan=2, pady=10)
        self.make_button(btn_frame, "🆕 جديد", self.clear_form, "#475569").pack(side="right", padx=4)
        self.make_button(btn_frame, "📚 تطبيق القالب", self.apply_template, WARN, "black").pack(side="right", padx=4)
        self.make_button(btn_frame, "💾 حفظ", self.save_note, SUCCESS).pack(side="right", padx=4)
        self.make_button(btn_frame, "✏️ تعديل", self.update_note, WARN, "black").pack(side="right", padx=4)
        self.make_button(btn_frame, "🗑️ حذف", self.delete_note, DANGER).pack(side="right", padx=4)
        self.make_button(btn_frame, "👁️ معاينة", self.preview_note, ACCENT).pack(side="right", padx=4)
        self.make_button(btn_frame, "🖨️ طباعة احترافية", self.print_note, "#0f766e").pack(side="right", padx=4)
        self.make_button(btn_frame, "📝 تصدير TXT", self.export_txt, "#7c3aed").pack(side="right", padx=4)
        self.make_button(btn_frame, "📄 إنشاء Word احترافي", self.export_word, "#4f46e5").pack(side="right", padx=4)

        top_search = tk.Frame(right, bg=BG)
        top_search.pack(fill="x", pady=(0, 8))
        tk.Label(top_search, text="بحث", bg=BG, fg=TEXT, font=("Arial", 11, "bold")).pack(side="right", padx=5)
        search_entry = tk.Entry(top_search, textvariable=self.search_var, justify="right", font=("Arial", 11), width=35, bg=INPUT_BG, fg=INPUT_FG)
        search_entry.pack(side="right", padx=5)
        search_entry.bind("<KeyRelease>", lambda e: self.refresh_table())
        self.make_button(top_search, "تحديث القائمة", self.refresh_table, ACCENT).pack(side="right", padx=4)

        table_frame = tk.LabelFrame(right, text="أرشيف المذكرات", font=("Arial", 12, "bold"), bg=PANEL, fg=GOLD)
        table_frame.pack(fill="both", expand=True)

        cols = ("teacher", "subject", "level", "lesson", "week", "session", "date")
        self.tree = ttk.Treeview(table_frame, columns=cols, show="headings", height=18)
        headers = {
            "teacher": "الأستاذ",
            "subject": "المادة",
            "level": "المستوى",
            "lesson": "عنوان الدرس",
            "week": "الأسبوع",
            "session": "الحصة",
            "date": "تاريخ الإنشاء",
        }
        widths = {"teacher": 150, "subject": 120, "level": 120, "lesson": 220, "week": 80, "session": 80, "date": 140}
        for col in cols:
            self.tree.heading(col, text=headers[col])
            self.tree.column(col, width=widths[col], anchor="center")
        self.tree.pack(fill="both", expand=True, padx=5, pady=5)
        self.tree.bind("<<TreeviewSelect>>", self.on_select)

        self.apply_template()
        self.update_current_labels()

    def update_current_labels(self):
        self.subject_var.set(self.entries["المادة"].get() or "-")
        self.level_var.set(self.entries["المستوى"].get() or "-")
        self.count_var.set(str(len(self.notes)))

    def apply_template(self):
        subject = self.entries["المادة"].get().strip()
        template = TEMPLATES.get(subject, {})
        for field in ["الكفاءة الختامية", "مؤشر الكفاءة", "الوسائل"]:
            self.entries[field].delete(0, tk.END)
            self.entries[field].insert(0, template.get(field, ""))
        for field in ["وضعية الانطلاق", "سير الحصة", "التقويم", "الوضعية الإدماجية", "الملاحظات"]:
            self.texts[field].delete("1.0", tk.END)
            self.texts[field].insert("1.0", template.get(field, ""))
        self.update_current_labels()

    def get_form_data(self):
        data = {label: widget.get().strip() for label, widget in self.entries.items()}
        for label, text in self.texts.items():
            data[label] = text.get("1.0", tk.END).strip()
        data["تاريخ الإنشاء"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        return data

    def validate_data(self, data):
        for field in ["اسم الأستاذ", "المادة", "المستوى", "عنوان الدرس"]:
            if not data.get(field):
                raise ValueError(f"الحقل مطلوب: {field}")

    def clear_form(self):
        self.current_index = None
        for label, widget in self.entries.items():
            if isinstance(widget, ttk.Combobox):
                widget.set("الأولى متوسط" if label == "المستوى" else "اللغة العربية")
            else:
                widget.delete(0, tk.END)
        for text in self.texts.values():
            text.delete("1.0", tk.END)
        self.apply_template()
        self.update_current_labels()

    def save_note(self):
        try:
            data = self.get_form_data()
            self.validate_data(data)
            self.notes.append(data)
            self.save_data()
            self.refresh_table()
            self.clear_form()
            messagebox.showinfo("نجاح", "تم حفظ المذكرة بنجاح")
        except ValueError as e:
            messagebox.showerror("خطأ", str(e))

    def update_note(self):
        if self.current_index is None:
            messagebox.showwarning("تنبيه", "اختر مذكرة من الجدول أولا")
            return
        try:
            original_date = self.notes[self.current_index].get("تاريخ الإنشاء", "")
            data = self.get_form_data()
            self.validate_data(data)
            data["تاريخ الإنشاء"] = original_date
            self.notes[self.current_index] = data
            self.save_data()
            self.refresh_table()
            messagebox.showinfo("نجاح", "تم تعديل المذكرة")
        except ValueError as e:
            messagebox.showerror("خطأ", str(e))

    def delete_note(self):
        if self.current_index is None:
            messagebox.showwarning("تنبيه", "اختر مذكرة من الجدول أولا")
            return
        if messagebox.askyesno("تأكيد", "هل تريد حذف هذه المذكرة؟"):
            del self.notes[self.current_index]
            self.save_data()
            self.refresh_table()
            self.clear_form()
            messagebox.showinfo("تم", "تم حذف المذكرة")

    def refresh_table(self):
        keyword = self.search_var.get().strip()
        for item in self.tree.get_children():
            self.tree.delete(item)
        for idx, note in enumerate(self.notes):
            text = " ".join(str(v) for v in note.values())
            if keyword and keyword not in text:
                continue
            self.tree.insert("", tk.END, iid=str(idx), values=(
                note.get("اسم الأستاذ", ""),
                note.get("المادة", ""),
                note.get("المستوى", ""),
                note.get("عنوان الدرس", ""),
                note.get("الأسبوع", ""),
                note.get("الحصة", ""),
                note.get("تاريخ الإنشاء", ""),
            ))
        self.count_var.set(str(len(self.notes)))

    def on_select(self, event=None):
        selection = self.tree.selection()
        if not selection:
            return
        self.current_index = int(selection[0])
        note = self.notes[self.current_index]
        for label, widget in self.entries.items():
            value = note.get(label, "")
            if isinstance(widget, ttk.Combobox):
                widget.set(value)
            else:
                widget.delete(0, tk.END)
                widget.insert(0, value)
        for label, text in self.texts.items():
            text.delete("1.0", tk.END)
            text.insert("1.0", note.get(label, ""))
        self.update_current_labels()

    def format_note(self, note):
        return f"""
==============================
        مذكرة تربوية
==============================
المؤسسة: {SCHOOL_NAME}
الولاية/المدينة: {SCHOOL_CITY}
الأستاذ: {note.get('اسم الأستاذ', '')}
المادة: {note.get('المادة', '')}
المستوى: {note.get('المستوى', '')}
المقطع: {note.get('المقطع', '')}
الوحدة: {note.get('الوحدة', '')}
عنوان الدرس: {note.get('عنوان الدرس', '')}
الأسبوع: {note.get('الأسبوع', '')}
الحصة: {note.get('الحصة', '')}
المدة: {note.get('المدة', '')}

الكفاءة الختامية:
{note.get('الكفاءة الختامية', '')}

مؤشر الكفاءة:
{note.get('مؤشر الكفاءة', '')}

الوسائل:
{note.get('الوسائل', '')}

وضعية الانطلاق:
{note.get('وضعية الانطلاق', '')}

سير الحصة:
{note.get('سير الحصة', '')}

التقويم:
{note.get('التقويم', '')}

الوضعية الإدماجية:
{note.get('الوضعية الإدماجية', '')}

الملاحظات:
{note.get('الملاحظات', '')}

تاريخ الإنشاء: {note.get('تاريخ الإنشاء', '')}
==============================
""".strip()

    def get_selected_or_current_note(self):
        if self.current_index is None:
            data = self.get_form_data()
            self.validate_data(data)
            return data
        return self.notes[self.current_index]

    def build_print_html(self, note):
        def esc(text):
            return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>")

        sections = [
            ("الكفاءة الختامية", note.get("الكفاءة الختامية", "")),
            ("مؤشر الكفاءة", note.get("مؤشر الكفاءة", "")),
            ("الوسائل", note.get("الوسائل", "")),
            ("وضعية الانطلاق", note.get("وضعية الانطلاق", "")),
            ("سير الحصة", note.get("سير الحصة", "")),
            ("التقويم", note.get("التقويم", "")),
            ("الوضعية الإدماجية", note.get("الوضعية الإدماجية", "")),
            ("الملاحظات", note.get("الملاحظات", "")),
        ]
        sections_html = "".join([
            f'<div class="section"><div class="section-title">{esc(title)}</div><div class="section-body">{esc(body)}</div></div>'
            for title, body in sections
        ])
        return f'''<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
<meta charset="utf-8">
<title>طباعة مذكرة</title>
<style>
body {{ font-family: Arial, sans-serif; background:#f3f4f6; margin:0; padding:24px; color:#111827; }}
.page {{ max-width:900px; margin:0 auto; background:white; padding:32px; border:2px solid #1e3a8a; box-shadow:0 8px 30px rgba(0,0,0,.08); }}
.header {{ text-align:center; border-bottom:3px solid #d4af37; padding-bottom:12px; margin-bottom:18px; }}
.header h1 {{ margin:0; color:#1e3a8a; font-size:28px; }}
.header p {{ margin:6px 0 0; color:#374151; font-size:15px; }}
.grid {{ display:grid; grid-template-columns:1fr 1fr; gap:10px 18px; margin-bottom:18px; }}
.card {{ background:#f9fafb; border:1px solid #e5e7eb; padding:10px 12px; border-radius:8px; }}
.label {{ font-weight:700; color:#1e3a8a; margin-left:6px; }}
.section {{ margin-top:14px; border:1px solid #d1d5db; border-radius:10px; overflow:hidden; }}
.section-title {{ background:#1e3a8a; color:white; padding:10px 14px; font-weight:700; }}
.section-body {{ padding:14px; line-height:1.9; min-height:36px; }}
.footer {{ margin-top:22px; text-align:center; color:#6b7280; font-size:13px; }}
@media print {{ body {{ background:white; padding:0; }} .page {{ box-shadow:none; border:none; max-width:none; padding:12mm; }} }}
</style>
</head>
<body>
<div class="page">
<div class="header">
<h1>مذكرة تربوية</h1>
<p>{esc(SCHOOL_NAME)} - {esc(SCHOOL_CITY)}</p>
</div>
<div class="grid">
<div class="card"><span class="label">الأستاذ:</span>{esc(note.get('اسم الأستاذ',''))}</div>
<div class="card"><span class="label">المادة:</span>{esc(note.get('المادة',''))}</div>
<div class="card"><span class="label">المستوى:</span>{esc(note.get('المستوى',''))}</div>
<div class="card"><span class="label">المقطع:</span>{esc(note.get('المقطع',''))}</div>
<div class="card"><span class="label">الوحدة:</span>{esc(note.get('الوحدة',''))}</div>
<div class="card"><span class="label">عنوان الدرس:</span>{esc(note.get('عنوان الدرس',''))}</div>
<div class="card"><span class="label">الأسبوع:</span>{esc(note.get('الأسبوع',''))}</div>
<div class="card"><span class="label">الحصة:</span>{esc(note.get('الحصة',''))}</div>
<div class="card"><span class="label">المدة:</span>{esc(note.get('المدة',''))}</div>
<div class="card"><span class="label">تاريخ الإنشاء:</span>{esc(note.get('تاريخ الإنشاء',''))}</div>
</div>
{sections_html}
<div class="footer">جاهزة للطباعة - استعمل أمر الطباعة من المتصفح أو حفظ PDF</div>
</div>
<script>
window.onload = function() {{
  setTimeout(function() {{ window.print(); }}, 300);
}};
</script>
</body>
</html>'''

    def preview_note(self):
        try:
            note = self.get_selected_or_current_note()
        except ValueError as e:
            messagebox.showerror("خطأ", str(e))
            return
        preview = tk.Toplevel(self.root)
        preview.title("معاينة المذكرة")
        preview.geometry("900x720")
        preview.configure(bg=BG)
        text = tk.Text(preview, wrap="word", font=("Arial", 12), bg=INPUT_BG, fg=INPUT_FG)
        text.pack(fill="both", expand=True, padx=10, pady=10)
        text.insert("1.0", self.format_note(note))

    def print_note(self):
        try:
            note = self.get_selected_or_current_note()
        except ValueError as e:
            messagebox.showerror("خطأ", str(e))
            return
        html = self.build_print_html(note)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".html", mode="w", encoding="utf-8") as tmp:
            tmp.write(html)
            temp_path = tmp.name
        webbrowser.open("file://" + os.path.abspath(temp_path))

    def export_txt(self):
        if self.current_index is None:
            messagebox.showwarning("تنبيه", "اختر مذكرة من الجدول أولا")
            return
        note = self.notes[self.current_index]
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")], initialfile="mudhakira.txt")
        if not path:
            return
        with open(path, "w", encoding="utf-8") as f:
            f.write(self.format_note(note))
        messagebox.showinfo("نجاح", "تم تصدير المذكرة")

    def export_word(self):
        try:
            note = self.get_selected_or_current_note()
        except ValueError as e:
            messagebox.showerror("خطأ", str(e))
            return

        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile="مذكرة_جاهزة_للطباعة.docx",
        )
        if not path:
            return

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("مذكرة تربوية جاهزة للطباعة")
        r.bold = True
        r.font.size = Pt(16)

        school = doc.add_paragraph()
        school.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s = school.add_run(f"{SCHOOL_NAME} - {SCHOOL_CITY}")
        s.bold = True
        s.font.size = Pt(12)

        doc.add_paragraph("")

        info_table = doc.add_table(rows=0, cols=4)
        info_table.style = "Table Grid"
        pairs = [
            ("اسم الأستاذ", note.get("اسم الأستاذ", ""), "المادة", note.get("المادة", "")),
            ("المستوى", note.get("المستوى", ""), "المقطع", note.get("المقطع", "")),
            ("الوحدة", note.get("الوحدة", ""), "عنوان الدرس", note.get("عنوان الدرس", "")),
            ("الأسبوع", note.get("الأسبوع", ""), "الحصة", note.get("الحصة", "")),
            ("المدة", note.get("المدة", ""), "تاريخ الإنشاء", note.get("تاريخ الإنشاء", "")),
        ]
        for a, b, c, d in pairs:
            row = info_table.add_row().cells
            row[0].text = a
            row[1].text = str(b)
            row[2].text = c
            row[3].text = str(d)

        sections = [
            "الكفاءة الختامية",
            "مؤشر الكفاءة",
            "الوسائل",
            "وضعية الانطلاق",
            "سير الحصة",
            "التقويم",
            "الوضعية الإدماجية",
            "الملاحظات",
        ]
        for title_text in sections:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rr = p.add_run(title_text)
            rr.bold = True
            rr.font.size = Pt(13)
            body = doc.add_paragraph(note.get(title_text, ""))
            body.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save(path)
        try:
            import subprocess
            if os.name == "nt":
                os.startfile(path)
            else:
                subprocess.call(["open", path])
        except Exception:
            pass
        messagebox.showinfo("نجاح", "تم إنشاء ملف Word احترافي وفتحه مباشرة")

    def save_data(self):
        with open(DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(self.notes, f, ensure_ascii=False, indent=2)

    def load_data(self):
        if os.path.exists(DATA_FILE):
            try:
                with open(DATA_FILE, "r", encoding="utf-8") as f:
                    self.notes = json.load(f)
            except Exception:
                self.notes = []


if __name__ == "__main__":
    root = tk.Tk()
    app = NotesApp(root)
    root.mainloop()
